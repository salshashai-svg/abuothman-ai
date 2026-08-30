import asyncio
import csv
import os
import threading
from pathlib import Path

from agents import Agent, Runner

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    import pandas as pd
except ImportError:  # pragma: no cover
    pd = None

from specialists.health_agent import health_agent
from specialists.finance_agent import finance_agent
from specialists.business_agent import business_agent
from specialists.travel_agent import travel_agent
from specialists.documents_agent import documents_agent
from specialists.communication_agent import communication_agent

chief_agent = Agent(
    name="AbuOthman Chief Agent",
    instructions="""
You are AbuOthman's Chief AI Agent, the intelligent orchestrator of the AbuOthman AI system.

Your primary responsibilities:
1. Route user requests to the most appropriate specialist agent
2. Ensure all document processing is handled by Documents Agent
3. Coordinate between specialists when multiple are needed
4. Provide coherent final answers by synthesizing specialist responses

Routing Rules:

**ALWAYS route to Documents Agent if:**
- The request mentions FILE_PATH= or DOCUMENT_CONTENT:
- An uploaded document is included in the message
- The user asks about document content, files, or attachments
- The request involves summarization, table extraction, or document analysis
- The user wants to extract information from files

**Route to Health Agent if:**
- Health conditions, medications, medical advice requested
- Questions about kidney disease, transplants, diabetes, blood pressure
- Laboratory test interpretation
- Medical recommendations needed

**Route to Finance Agent if:**
- Financial planning or investment questions
- Budget, accounting, or financial analysis
- Tax or economic questions
- Business financial matters

**Route to Business Agent if:**
- Business strategy or development questions
- Company operations or management
- Business analysis or planning
- Entrepreneurship questions

**Route to Travel Agent if:**
- Travel planning, destinations, or itineraries
- Tourism information or recommendations
- Travel logistics or bookings

**Route to Communication Agent if:**
- Writing, editing, or communication help
- Professional communication or presentation
- Language improvement or translation
- Content creation

Important Guidelines:
- ALWAYS prioritize Documents Agent for any file-related requests
- Use handoffs to route to specialists - never answer specialist questions yourself
- When multiple specialists are needed, coordinate them sequentially and synthesize results
- Maintain context from DOCUMENT_CONTENT when routing to other agents
- Provide clear, accurate, and professional responses
- If uncertain about routing, ask the user for clarification

Handoff Coordination:
- If Documents Agent extracts data, pass it to other specialists if further analysis is needed
- Keep all specialist responses focused on their domain
- Always return a final, unified answer that addresses the user's complete request
""",
    handoffs=[
        health_agent,
        finance_agent,
        business_agent,
        travel_agent,
        documents_agent,
        communication_agent,
    ],
)


def extract_text_from_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as file:
            return file.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="latin-1", errors="replace") as file:
            return file.read()


def extract_text_from_csv(path: str) -> str:
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace", newline="") as file:
        reader = csv.reader(file)
        for row in reader:
            rows.append(", ".join(row))
    return "\n".join(rows)


def extract_text_from_pdf(path: str) -> str:
    if PdfReader is None:
        return "PDF extraction library is not available."

    pages = []
    try:
        reader = PdfReader(path)
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages) if pages else "No readable text was found in this PDF."
    except Exception as exc:  # pragma: no cover
        return f"Unable to read PDF content: {exc}"


def extract_text_from_docx(path: str) -> str:
    if Document is None:
        return "DOCX extraction library is not available."

    try:
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(cells))
        return "\n".join(paragraphs + table_text)
    except Exception as exc:  # pragma: no cover
        return f"Unable to read Word document: {exc}"


def extract_text_from_excel(path: str) -> str:
    if pd is None:
        return "Excel extraction library is not available."

    try:
        dataframe = pd.read_excel(path)
        return dataframe.to_string(index=False)
    except Exception as exc:  # pragma: no cover
        return f"Unable to read Excel file: {exc}"


def extract_text_from_image(path: str) -> str:
    if Image is None or pytesseract is None:
        return "OCR libraries are not installed. Image content could not be read."

    try:
        image = Image.open(path)
        text = pytesseract.image_to_string(image)
        return text.strip() or "No readable text was detected in this image."
    except Exception as exc:  # pragma: no cover
        return f"Unable to read image content: {exc}"


def extract_document_content(uploaded_file) -> str:
    if hasattr(uploaded_file, "read"):
        path = getattr(uploaded_file, "name", None)
        if path:
            uploaded_path = Path(path)
            if uploaded_path.exists():
                file_path = str(uploaded_path)
            else:
                temp_dir = Path("/tmp")
                temp_dir.mkdir(exist_ok=True)
                temp_file = temp_dir / f"uploaded_{abs(hash(uploaded_file.name))}.bin"
                temp_file.write_bytes(uploaded_file.getvalue())
                file_path = str(temp_file)
        else:
            file_path = None
    else:
        file_path = str(uploaded_file)

    if not file_path or not os.path.exists(file_path):
        return "No readable content could be extracted because the uploaded file was not found."

    extension = Path(file_path).suffix.lower()
    file_handlers = {
        ".txt": extract_text_from_txt,
        ".csv": extract_text_from_csv,
        ".pdf": extract_text_from_pdf,
        ".doc": extract_text_from_docx,
        ".docx": extract_text_from_docx,
        ".xls": extract_text_from_excel,
        ".xlsx": extract_text_from_excel,
        ".png": extract_text_from_image,
        ".jpg": extract_text_from_image,
        ".jpeg": extract_text_from_image,
        ".webp": extract_text_from_image,
        ".bmp": extract_text_from_image,
    }

    handler = file_handlers.get(extension)
    if handler is None:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as file:
                return file.read()
        except Exception:
            return "This file type is not supported for content extraction."

    return handler(file_path)


async def ask_agent_async(message: str) -> str:
    result = await Runner.run(
        chief_agent,
        message,
    )
    return result.final_output


def ask_agent(message: str, uploaded_file=None) -> str:
    if uploaded_file:
        document_text = extract_document_content(uploaded_file)
        file_label = uploaded_file if isinstance(uploaded_file, str) else getattr(uploaded_file, "name", "uploaded_document")
        message = (
            f"{message}\n\n"
            f"FILE_PATH={file_label}\n\n"
            f"DOCUMENT_CONTENT:\n{document_text}"
        )

    try:
        asyncio.get_running_loop()
        result = {}

        def runner():
            result["value"] = asyncio.run(ask_agent_async(message))

        thread = threading.Thread(target=runner)
        thread.start()
        thread.join()
        return result.get("value", "")
    except RuntimeError:
        return asyncio.run(ask_agent_async(message))